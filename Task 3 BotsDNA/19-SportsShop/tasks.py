from robocorp.tasks import task
from robocorp import browser
import openpyxl
from openpyxl import Workbook
import csv
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
import time
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.relative_locator import locate_with
from selenium.webdriver.support.select import Select
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import PyPDF2
from docx.shared import Inches
from PIL import Image
import io
import requests

@task
def sports_shop():
    """Get details of the latest sports equipment and send emails about it to the schools"""
    driver = webdriver.Chrome()
    driver.get("https://botsdna.com/sportshop/")
    
    wait = WebDriverWait(driver, 10)
    wait.until(EC.presence_of_element_located((By.XPATH, "//table[.//span[contains(@style, 'color:red')]]")))
    
    offer=0
    # Locate the table containing the red color span element
    new_stock = driver.find_element(By.XPATH, ".//span[contains(@style, 'color:red')]")
    tbody = driver.find_element(locate_with(By.TAG_NAME, "tbody").above(new_stock))
    details = driver.find_element(locate_with(By.TAG_NAME,"tbody").below(tbody))
    
    #['Skipping Rope', 'Product Name : Skipping Rope', 'Product Code : DG18189', 'Price : 500/-', 'Stock Arrived From : Kochi']
    data = details.text.strip().split('\n')
    cost = data[3]
    if 'Price' in cost:
        price = int(cost.split(':')[1].strip().split('/')[0])
    image = details.find_element(By.TAG_NAME,"img")
    src = image.get_attribute('src')
    response = requests.get(src)
    img = Image.open(io.BytesIO(response.content))
    img.save('image.jpg')

    driver.get("https://botsdna.com/sportshop/schools.html")
    rows = driver.find_elements(By.XPATH,"//*[@id='courts']/tbody/tr")
    for i in range(1,len(rows)):
        row_data = rows[i].find_elements(By.TAG_NAME,"td")
        school_code = row_data[0].text
        school_name = row_data[1].text
        student_strength = int(row_data[2].text)
        if student_strength < 200:
            offer = 10
        elif student_strength > 200 and student_strength<=600:
            offer = 15
        elif student_strength > 600 and student_strength<=1000:
            offer = 20
        else:
            offer = 25
        price = (((100-offer)/100)*price)
        workbook = openpyxl.load_workbook('EmailsDatabase.xlsx')
        sheet = workbook.get_sheet_by_name("SchoolsData")

        for row in sheet.iter_rows():
            # Check if the first column of the row contains the search value
            if row[0].value == school_code:
                # If the search value is found, print the row number
                row_value =  row[0].row
                break
        if sheet.cell(row = row_value, column=2).value != "":            
            email_id = sheet.cell(row = row_value, column=1).value
        elif sheet.cell(row = row_value, column=3).value != "":
            email_id = sheet.cell(row = row_value, column=2).value
        elif sheet.cell(row = row_value, column=4).value != "":
            email_id = sheet.cell(row = row_value, column=4).value
        else:
            email_id = sheet.cell(row = row_value, column=5).value

        doc = Document('SportsTemplet.docx')

        # Add a paragraph with the given text
        doc.add_paragraph(f'Your School {school_name} Has {offer}% Discount on All Products')

        # Add a paragraph with the product name
        doc.add_paragraph(data[1])

        # Add a paragraph with the product code
        doc.add_paragraph(data[2])

        # Add a paragraph with the price
        doc.add_paragraph(f'Unit Price: {cost}, Discounted price: {price}')

        # Add the saved image to the document
        doc.add_picture('image.jpg', width=Inches(3))

        # Save the Word document
        doc.save(f'{school_name}.docx')


    
    
